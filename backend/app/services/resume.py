import json
import logging
import uuid

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import joinedload

from app.core.exceptions import NotFoundError
from app.models import ResumeSection, ResumeVersion
from app.repositories import (
    ResumeSectionRepository,
    ResumeVersionRepository,
)
from app.schemas.resume import ResumeImportData
from app.services.audit import AuditService
from database.models.career_profile import CareerProfile

logger = logging.getLogger(__name__)

# Canonical resume origins. "master" is the legacy default for manually
# created resumes; the library surfaces the others as origin badges.
ORIGIN_MASTER = "master"
ORIGIN_UPLOADED = "uploaded"
ORIGIN_AI_GENERATED = "ai_generated"
ORIGIN_AI_TAILORED = "ai_tailored"

ORIGIN_MASTER_RESUMES = (ORIGIN_MASTER, ORIGIN_UPLOADED)
ORIGIN_AI_RESUMES = (ORIGIN_AI_GENERATED, ORIGIN_AI_TAILORED)


class ResumeService:
    def __init__(self, session: AsyncSession):
        self.session = session
        self.resume_repo = ResumeVersionRepository(session)
        self.section_repo = ResumeSectionRepository(session)
        self.audit_service = AuditService(session)

    # ── Resume CRUD ──

    async def list_resumes(
        self, user_id: uuid.UUID, archived: bool | None = None, origin: list[str] | None = None
    ) -> list[ResumeVersion]:
        if origin:
            # The library "master" tab shows manual + uploaded resumes.
            if set(origin) == {ORIGIN_MASTER}:
                origin = list(ORIGIN_MASTER_RESUMES)
            return await self.resume_repo.list_by_user_and_origins(user_id, origin)
        return await self.resume_repo.list_by_user(user_id, archived=archived)

    async def get_resume(self, resume_id: uuid.UUID, user_id: uuid.UUID) -> ResumeVersion:
        resume = await self.resume_repo.get_with_sections(resume_id)
        if not resume or resume.user_id != user_id:
            raise NotFoundError("Resume not found.")
        return resume

    async def create_resume(
        self,
        user_id: uuid.UUID,
        title: str | None,
        description: str | None = None,
        template: str | None = None,
        resume_type: str | None = None,
        change_summary: str | None = None,
        sections: list[dict] | None = None,
        origin: str = ORIGIN_MASTER,
    ) -> ResumeVersion:
        latest_version = await self.resume_repo.latest_version(user_id)
        resume = ResumeVersion(
            user_id=user_id,
            version=latest_version + 1,
            title=title or f"Resume v{latest_version + 1}",
            description=description,
            template=template,
            resume_type=resume_type,
            change_summary=change_summary,
            is_default=latest_version == 0,
            origin=origin,
            status="active" if sections else "draft",
        )
        created = await self.resume_repo.create(resume)

        if sections:
            for i, section_data in enumerate(sections):
                section = ResumeSection(
                    resume_id=created.id,
                    section_type=section_data.get("section_type", "custom"),
                    title=section_data.get("title"),
                    content=section_data.get("content"),
                    sort_order=section_data.get("sort_order", i),
                    visible=section_data.get("visible", True),
                )
                await self.section_repo.create(section)

        await self.audit_service.log(
            "RESUME_CREATED",
            user_id=user_id,
            entity="resume",
            entity_id=created.id,
            outcome="success",
        )
        return await self.resume_repo.get_with_sections(created.id)

    async def update_resume(
        self,
        resume_id: uuid.UUID,
        user_id: uuid.UUID,
        data: dict,
    ) -> ResumeVersion:
        resume = await self.resume_repo.get_by_id(resume_id)
        if not resume or resume.user_id != user_id:
            raise NotFoundError("Resume not found.")
        for key, value in data.items():
            if value is not None and hasattr(resume, key):
                setattr(resume, key, value)
        await self.resume_repo.update(resume)
        await self.audit_service.log(
            "RESUME_UPDATED",
            user_id=user_id,
            entity="resume",
            entity_id=resume_id,
            outcome="success",
        )
        return await self.resume_repo.get_with_sections(resume_id)

    async def delete_resume(self, resume_id: uuid.UUID, user_id: uuid.UUID) -> None:
        resume = await self.resume_repo.get_by_id(resume_id)
        if not resume or resume.user_id != user_id:
            raise NotFoundError("Resume not found.")
        await self.resume_repo.delete(resume)
        await self.audit_service.log(
            "RESUME_DELETED",
            user_id=user_id,
            entity="resume",
            entity_id=resume_id,
            outcome="success",
        )

    # ── Status Management ──

    async def archive_resume(self, resume_id: uuid.UUID, user_id: uuid.UUID) -> ResumeVersion:
        resume = await self.resume_repo.get_by_id(resume_id)
        if not resume or resume.user_id != user_id:
            raise NotFoundError("Resume not found.")
        resume.status = "archived"
        resume.archived = True
        await self.resume_repo.update(resume)
        await self.audit_service.log(
            "RESUME_ARCHIVED",
            user_id=user_id,
            entity="resume",
            entity_id=resume_id,
            outcome="success",
        )
        return resume

    async def restore_resume(self, resume_id: uuid.UUID, user_id: uuid.UUID) -> ResumeVersion:
        resume = await self.resume_repo.get_by_id(resume_id)
        if not resume or resume.user_id != user_id:
            raise NotFoundError("Resume not found.")
        resume.status = "active"
        resume.archived = False
        await self.resume_repo.update(resume)
        await self.audit_service.log(
            "RESUME_RESTORED",
            user_id=user_id,
            entity="resume",
            entity_id=resume_id,
            outcome="success",
        )
        return resume

    async def set_default_resume(self, resume_id: uuid.UUID, user_id: uuid.UUID) -> ResumeVersion:
        resume = await self.resume_repo.get_by_id(resume_id)
        if not resume or resume.user_id != user_id:
            raise NotFoundError("Resume not found.")
        return await self.resume_repo.set_default(resume_id, user_id)

    # ── Sections ──

    async def add_section(
        self,
        resume_id: uuid.UUID,
        user_id: uuid.UUID,
        section_data: dict,
    ) -> ResumeSection:
        resume = await self.resume_repo.get_by_id(resume_id)
        if not resume or resume.user_id != user_id:
            raise NotFoundError("Resume not found.")
        section = ResumeSection(
            resume_id=resume_id,
            section_type=section_data.get("section_type", "custom"),
            title=section_data.get("title"),
            content=section_data.get("content"),
            sort_order=section_data.get("sort_order", 0),
            visible=section_data.get("visible", True),
        )
        return await self.section_repo.create(section)

    async def update_section(
        self,
        section_id: uuid.UUID,
        user_id: uuid.UUID,
        data: dict,
    ) -> ResumeSection:
        section = await self.section_repo.get_by_id(section_id)
        if not section:
            raise NotFoundError("Section not found.")
        resume = await self.resume_repo.get_by_id(section.resume_id)
        if not resume or resume.user_id != user_id:
            raise NotFoundError("Section not found.")
        for key, value in data.items():
            if value is not None and hasattr(section, key):
                setattr(section, key, value)
        await self.section_repo.update(section)
        return section

    async def delete_section(self, section_id: uuid.UUID, user_id: uuid.UUID) -> None:
        section = await self.section_repo.get_by_id(section_id)
        if not section:
            raise NotFoundError("Section not found.")
        resume = await self.resume_repo.get_by_id(section.resume_id)
        if not resume or resume.user_id != user_id:
            raise NotFoundError("Section not found.")
        await self.section_repo.delete(section)

    # ── Import / Export ──

    async def import_resume(self, user_id: uuid.UUID, data: ResumeImportData) -> ResumeVersion:
        sections_data = [s.model_dump() for s in data.sections]
        return await self.create_resume(
            user_id=user_id,
            title=data.title,
            description=data.description,
            template=data.template,
            resume_type=data.resume_type,
            change_summary=data.change_summary or "Imported resume",
            sections=sections_data,
        )

    async def export_resume(self, resume_id: uuid.UUID, user_id: uuid.UUID) -> ResumeVersion:
        return await self.get_resume(resume_id, user_id)

    # ── Upload / Parse ──

    RESUME_HEADINGS = [
        "summary", "professional summary", "profile", "about me",
        "experience", "work experience", "employment", "professional experience",
        "education", "academic background", "academic",
        "skills", "technical skills", "core competencies", "technologies",
        "projects", "personal projects", "professional projects",
        "certifications", "certificates", "licenses",
        "languages", "language proficiency",
        "publications", "research",
        "awards", "honors", "achievements",
        "links", "social links", "profiles",
        "additional", "additional sections", "other",
    ]

    def _classify_section(self, heading: str) -> str:
        h = heading.lower().strip().rstrip(":")
        for category, keywords in {
            "summary": ["summary", "professional summary", "profile", "about me", "objective", "career objective"],
            "experience": ["experience", "work experience", "employment", "professional experience", "work history"],
            "education": ["education", "academic", "university", "college", "school", "degree", "qualification"],
            "skills": ["skills", "competencies", "technologies", "technical", "tools", "expertise"],
            "projects": ["projects", "side projects", "professional projects"],
            "certifications": ["certifications", "certificates", "licenses", "professional development"],
            "languages": ["languages", "language"],
            "publications": ["publications", "research", "papers"],
            "awards": ["awards", "honors", "achievements", "recognition"],
            "links": ["links", "profiles", "social", "github", "linkedin", "portfolio"],
        }.items():
            if any(kw in h for kw in keywords):
                return category
        return "custom"

    def _parse_pdf_text(self, content: bytes) -> list[dict]:
        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
        except Exception:
            return [{"section_type": "custom", "title": "Extracted Content", "content": {"text": "Could not parse PDF."}, "sort_order": 0}]

        return self._segment_text(text)

    def _parse_docx_text(self, content: bytes) -> list[dict]:
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            paragraphs = [(p.text, p.style.name if p.style else "") for p in doc.paragraphs]
        except Exception:
            return [{"section_type": "custom", "title": "Extracted Content", "content": {"text": "Could not parse DOCX."}, "sort_order": 0}]

        sections = []
        current_type = "summary"
        current_title = "Professional Summary"
        current_lines = []

        for text, style in paragraphs:
            if not text.strip():
                continue
            is_heading = any(h in style.lower() for h in ["heading", "title", "header"]) or (
                text.strip().rstrip(":").lower() in self.RESUME_HEADINGS
                or any(h in text.strip().rstrip(":").lower() for h in self.RESUME_HEADINGS)
            )
            if is_heading:
                if current_lines:
                    sections.append({
                        "section_type": current_type,
                        "title": current_title,
                        "content": {"text": "\n".join(current_lines).strip()},
                    })
                heading_text = text.strip().rstrip(":")
                current_type = self._classify_section(heading_text)
                current_title = heading_text
                current_lines = []
            elif current_type:
                current_lines.append(text.strip())

        if current_lines:
            sections.append({
                "section_type": current_type,
                "title": current_title,
                "content": {"text": "\n".join(current_lines).strip()},
            })

        return [s for s in sections if s.get("content", {}).get("text")]

    def _segment_text(self, text: str) -> list[dict]:
        lines = text.split("\n")
        sections = []
        current_type = "summary"
        current_title = "Professional Summary"
        current_lines = []

        def is_section_heading(line: str) -> bool:
            clean = line.strip().rstrip(":").lower()
            if clean in self.RESUME_HEADINGS:
                return True
            if sum(1 for h in self.RESUME_HEADINGS if h in clean) >= 1 and len(clean) < 60:
                return True
            return False

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if is_section_heading(stripped):
                if current_lines:
                    sections.append({
                        "section_type": current_type,
                        "title": current_title,
                        "content": {"text": "\n".join(current_lines).strip()},
                    })
                heading_text = stripped.rstrip(":")
                current_type = self._classify_section(heading_text)
                current_title = heading_text
                current_lines = []
            else:
                current_lines.append(stripped)

        if current_lines:
            sections.append({
                "section_type": current_type,
                "title": current_title,
                "content": {"text": "\n".join(current_lines).strip()},
            })

        return [s for s in sections if s.get("content", {}).get("text")]

    async def parse_upload(self, content: bytes, filename: str) -> dict:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        title = filename.rsplit(".", 1)[0] if "." in filename else filename

        try:
            if ext == "pdf":
                sections = self._parse_pdf_text(content)
            elif ext == "docx":
                sections = self._parse_docx_text(content)
            else:
                text = content.decode("utf-8", errors="replace")
                sections = self._segment_text(text)

            confidence = min(90, 50 + len(sections) * 10) if sections else 0

            needs_review = []
            known_types = {"summary", "experience", "education", "skills", "projects", "certifications", "languages", "achievements", "publications", "awards", "links"}
            for s in sections:
                if s["section_type"] not in known_types or not s.get("content", {}).get("text"):
                    needs_review.append(s["section_type"])
                elif s["section_type"] in ("custom",) and not s.get("title"):
                    needs_review.append(s["section_type"])

            if not sections:
                sections = [{
                    "section_type": "summary",
                    "title": "Professional Summary",
                    "content": {"text": "Uploaded resume. Please review extracted content."},
                }]
                confidence = 30
                needs_review = ["summary"]

            for i, s in enumerate(sections):
                s["sort_order"] = i

            return {
                "title": title,
                "sections": sections,
                "confidence": confidence,
                "needs_review": needs_review,
            }
        except Exception:
            return {
                "title": title,
                "sections": [{
                    "section_type": "summary",
                    "title": "Professional Summary",
                    "content": {"text": "Uploaded resume. Please review and update extracted content."},
                    "sort_order": 0,
                }],
                "confidence": 20,
                "needs_review": ["summary"],
            }

    async def reorder_sections(
        self,
        resume_id: uuid.UUID,
        user_id: uuid.UUID,
        order: list[dict],
    ) -> list:
        resume = await self.resume_repo.get_with_sections(resume_id)
        if not resume or resume.user_id != user_id:
            raise NotFoundError("Resume not found.")
        updated = []
        for item in order:
            section_id = item.get("section_id")
            sort_order = item.get("sort_order", 0)
            for section in (resume.sections or []):
                if str(section.id) == str(section_id):
                    section.sort_order = sort_order
                    await self.section_repo.update(section)
                    updated.append(section)
                    break
        return updated

    # ── Generate From Profile ──

    async def generate_from_profile(
        self,
        user_id: uuid.UUID,
        title: str = "Generated Resume",
        template: str | None = None,
        section_filter: list[str] | None = None,
        enhance_with_ai: bool = False,
    ) -> ResumeVersion:
        sections_data = []
        stmt = select(CareerProfile).options(
            joinedload(CareerProfile.education),
            joinedload(CareerProfile.experience),
            joinedload(CareerProfile.projects),
            joinedload(CareerProfile.skills),
            joinedload(CareerProfile.certifications),
            joinedload(CareerProfile.languages),
            joinedload(CareerProfile.achievements),
        ).where(CareerProfile.user_id == user_id)
        result = await self.session.execute(stmt)
        profile = result.unique().scalar_one_or_none()

        want = section_filter or ["summary", "experience", "education", "skills", "projects"]

        if "summary" in want:
            headline = (profile.headline if profile else None) or ""
            bio = (profile.professional_summary if profile else None) or ""
            text = f"{headline}\n\n{bio}" if bio else headline or "Professional with experience in the field."
            sections_data.append({
                "section_type": "summary",
                "title": "Professional Summary",
                "content": {"text": text},
                "sort_order": 0,
            })

        if "experience" in want and profile and profile.experience:
            exp_text = "\n\n".join(
                f"{e.title} at {e.company}\n{getattr(e, 'description', '') or ''}"
                for e in profile.experience
            )
            sections_data.append({
                "section_type": "experience",
                "title": "Experience",
                "content": {"text": exp_text},
                "sort_order": 1,
            })

        if "education" in want and profile and profile.education:
            edu_text = "\n\n".join(
                f"{e.degree} - {e.institution}\n{getattr(e, 'field_of_study', '') or ''}"
                for e in profile.education
            )
            sections_data.append({
                "section_type": "education",
                "title": "Education",
                "content": {"text": edu_text},
                "sort_order": 2,
            })

        if "skills" in want and profile and profile.skills:
            skill_text = ", ".join(s.name for s in profile.skills)
            sections_data.append({
                "section_type": "skills",
                "title": "Skills",
                "content": {"text": skill_text},
                "sort_order": 3,
            })

        if "projects" in want and profile and profile.projects:
            proj_text = "\n\n".join(
                f"{p.name}\n{getattr(p, 'description', '') or ''}"
                for p in profile.projects
            )
            sections_data.append({
                "section_type": "projects",
                "title": "Projects",
                "content": {"text": proj_text},
                "sort_order": 4,
            })

        if "certifications" in want and profile and profile.certifications:
            cert_text = "\n".join(
                f"{c.name}" + (f" - {c.issuer}" if getattr(c, "issuer", None) else "")
                for c in profile.certifications
            )
            sections_data.append({
                "section_type": "certifications",
                "title": "Certifications",
                "content": {"text": cert_text},
                "sort_order": 5,
            })

        if "languages" in want and profile and profile.languages:
            lang_text = ", ".join(
                f"{lang.language}" + (f" ({lang.proficiency})" if getattr(lang, "proficiency", None) else "")
                for lang in profile.languages
            )
            sections_data.append({
                "section_type": "languages",
                "title": "Languages",
                "content": {"text": lang_text},
                "sort_order": 6,
            })

        if "achievements" in want and profile and profile.achievements:
            ach_text = "\n".join(
                f"{a.title}" + (f" - {a.organization}" if getattr(a, "organization", None) else "")
                for a in profile.achievements
            )
            sections_data.append({
                "section_type": "achievements",
                "title": "Achievements",
                "content": {"text": ach_text},
                "sort_order": 7,
            })

        if enhance_with_ai:
            sections_data = await self._ai_enhance_sections(sections_data, user_id)

        return await self.create_resume(
            user_id=user_id,
            title=title,
            template=template,
            resume_type="generated",
            change_summary="Generated from career profile" + (" with AI enhancement" if enhance_with_ai else ""),
            sections=sections_data,
            origin=ORIGIN_AI_GENERATED,
        )

    @staticmethod
    def _extract_improved_text(result: dict | None, current: str) -> str:
        """Best-effort parse of the AI improve response into plain text.

        The resume-improvement prompt asks for JSON {improved_bullets,
        summary, changes_made}. Some providers return clean JSON, others
        wrap it in prose; fall back to the raw content, then to the
        original content, so a bad response never corrupts the resume.
        """
        if not result:
            return current
        raw = str(result.get("improved_content", "") or "").strip()
        if not raw:
            return current
        try:
            parsed = json.loads(raw)
            if isinstance(parsed, dict):
                bullets = parsed.get("improved_bullets") or parsed.get("bullet_points")
                summary = parsed.get("summary")
                if isinstance(bullets, list) and len(bullets) > 0:
                    lines = [str(b) for b in bullets]
                    if summary:
                        lines.insert(0, str(summary))
                    return "\n".join(l for l in lines if l)
                if summary:
                    return str(summary)
        except (json.JSONDecodeError, ValueError):
            pass
        return raw

    async def _ai_enhance_sections(self, sections_data: list[dict], user_id: uuid.UUID) -> list[dict]:
        from app.ai.features.resume import ai_improve_resume_section

        enhanced = []
        for section in sections_data:
            section_type = section["section_type"]
            content_text = section.get("content", {}).get("text", "")
            if not content_text.strip():
                enhanced.append(section)
                continue
            try:
                result = await ai_improve_resume_section(
                    section_type=section_type,
                    current_content=content_text,
                    improvement_areas="grammar, tone, action_verbs, keywords",
                )
                improved = self._extract_improved_text(result, content_text)
                if improved and improved.strip():
                    section["content"]["text"] = improved
            except Exception:
                logger.exception("AI section enhancement failed; keeping original content.")
            enhanced.append(section)
        return enhanced

    # ── Optimize For Job ──

    async def optimize_for_job(
        self,
        resume_id: uuid.UUID,
        user_id: uuid.UUID,
        job_id: uuid.UUID,
        target_role: str | None = None,
        enhance_with_ai: bool = False,
    ) -> ResumeVersion:
        resume = await self.resume_repo.get_with_sections(resume_id)
        if not resume or resume.user_id != user_id:
            raise NotFoundError("Resume not found.")
        latest_version = await self.resume_repo.latest_version(user_id)
        new_resume = ResumeVersion(
            user_id=user_id,
            version=latest_version + 1,
            title=resume.title,
            description=resume.description,
            template=resume.template,
            resume_type="optimized",
            source="optimization",
            status="active",
            change_summary=f"Optimized for job {job_id}" + (f" ({target_role})" if target_role else "") + (" with AI enhancement" if enhance_with_ai else ""),
            previous_version_id=resume.id,
            generated_for_job_id=job_id,
            origin=ORIGIN_AI_TAILORED,
        )
        created = await self.resume_repo.create(new_resume)
        for section in (resume.sections or []):
            if enhance_with_ai:
                try:
                    from app.ai.features.resume import ai_improve_resume_section

                    original = section.content
                    content_text = (
                        original.get("text", "")
                        if isinstance(original, dict)
                        else str(original)
                        if original
                        else ""
                    )
                    result = None
                    if content_text.strip():
                        result = await ai_improve_resume_section(
                            section_type=section.section_type,
                            current_content=content_text,
                            target_role=target_role or "",
                            job_context=f"Optimizing for job {job_id}",
                            improvement_areas="grammar, tone, action_verbs, keywords, ats_optimization",
                        )
                    improved_text = self._extract_improved_text(result, content_text)
                    if isinstance(original, dict):
                        section_content = dict(original)
                        section_content["text"] = improved_text
                    else:
                        section_content = {"text": improved_text}
                except Exception:
                    logger.exception("AI optimization failed for section; copying original content.")
                    section_content = section.content
            else:
                section_content = section.content

            new_section = ResumeSection(
                resume_id=created.id,
                section_type=section.section_type,
                title=section.title,
                content=section_content,
                sort_order=section.sort_order,
                visible=section.visible,
            )
            await self.section_repo.create(new_section)
        await self.audit_service.log(
            "RESUME_OPTIMIZED",
            user_id=user_id,
            entity="resume",
            entity_id=created.id,
            outcome="success",
            details={"job_id": str(job_id), "enhanced_with_ai": enhance_with_ai},
        )
        return await self.resume_repo.get_with_sections(created.id)

    # ── ATS Analysis ──

    KNOWN_TECH_KEYWORDS = {
        "python", "javascript", "typescript", "java", "c++", "go", "rust", "sql",
        "react", "angular", "vue", "node", "django", "fastapi", "flask", "spring",
        "docker", "kubernetes", "aws", "azure", "gcp", "terraform", "ansible",
        "git", "ci/cd", "jenkins", "github actions", "linux", "rest api", "graphql",
        "redis", "postgresql", "mongodb", "mysql", "elasticsearch", "kafka",
        "machine learning", "deep learning", "nlp", "computer vision", "llm",
        "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy",
        "agile", "scrum", "jira", "confluence",
        "leadership", "communication", "team management", "mentoring",
    }

    SECTION_WEIGHTS = {
        "summary": 0.10,
        "experience": 0.35,
        "education": 0.15,
        "skills": 0.20,
        "projects": 0.10,
        "certifications": 0.05,
        "languages": 0.03,
        "publications": 0.02,
    }

    def _extract_text(self, sections: list) -> str:
        texts = []
        for s in (sections or []):
            content = s.content if hasattr(s, "content") else s.get("content", {})
            if isinstance(content, dict):
                texts.append(content.get("text", "") or "")
            elif isinstance(content, str):
                texts.append(content)
        return "\n".join(texts)

    def _count_keywords(self, text: str) -> dict:
        lower = text.lower()
        found = {}
        for kw in self.KNOWN_TECH_KEYWORDS:
            count = lower.count(kw)
            if count > 0:
                found[kw] = count
        return found

    async def analyze_ats(self, resume_id: uuid.UUID, user_id: uuid.UUID, job_id: uuid.UUID | None = None) -> dict:
        resume = await self.resume_repo.get_with_sections(resume_id)
        if not resume or resume.user_id != user_id:
            raise NotFoundError("Resume not found.")
        sections = resume.sections or []
        text = self._extract_text(sections)

        keywords = self._count_keywords(text)
        keyword_count = len(keywords)

        section_types = {s.section_type for s in sections}
        has_contact = any("contact" in (s.section_type) or "link" in (s.section_type) for s in sections)
        has_summary = "summary" in section_types or any("summary" in (s.title or "").lower() for s in sections)
        has_experience = "experience" in section_types
        has_education = "education" in section_types
        has_skills = "skills" in section_types

        keyword_score = min(100, keyword_count * 10 + 20)
        skills_score = 100 if has_skills else 40
        experience_score = 100 if has_experience else 30
        education_score = 100 if has_education else 40
        formatting_score = 85 if has_summary else 65
        sections_score = min(100, len(section_types) * 12 + 30)
        contact_score = 100 if has_contact else 60

        overall = round(
            keyword_score * 0.25
            + skills_score * 0.20
            + experience_score * 0.20
            + education_score * 0.10
            + formatting_score * 0.10
            + sections_score * 0.10
            + contact_score * 0.05
        )

        all_keywords = set(self.KNOWN_TECH_KEYWORDS)

        categories = [
            {
                "name": "Keywords",
                "score": keyword_score,
                "reason": f"Found {keyword_count} relevant keywords in your resume." if keyword_count > 0 else "No relevant keywords detected.",
                "suggestion": "Include relevant technologies naturally within your experience and skills sections.",
                "missing": sorted(list(all_keywords - set(keywords.keys())))[:8] if keyword_count < 10 else [],
            },
            {
                "name": "Skills",
                "score": skills_score,
                "reason": "Skills section found and populated." if has_skills else "No dedicated skills section detected.",
                "suggestion": "Create a skills section listing technical proficiencies.",
                "missing": [],
            },
            {
                "name": "Experience",
                "score": experience_score,
                "reason": "Experience section found." if has_experience else "No experience section detected.",
                "suggestion": "Add work experience with detailed descriptions and achievements.",
                "missing": [],
            },
            {
                "name": "Education",
                "score": education_score,
                "reason": "Education section found." if has_education else "No education section detected.",
                "suggestion": "Add your educational background including degrees and institutions.",
                "missing": [],
            },
            {
                "name": "Formatting",
                "score": formatting_score,
                "reason": "Good structure with summary section." if has_summary else "Missing professional summary.",
                "suggestion": "Add a professional summary at the top for better ATS parsing.",
                "missing": [],
            },
            {
                "name": "Sections",
                "score": sections_score,
                "reason": f"Contains {len(section_types)} section types." if section_types else "No sections detected.",
                "suggestion": "Include at least 5-6 different section types for completeness.",
                "missing": [s for s in ["summary", "experience", "education", "skills", "projects", "certifications"] if s not in section_types],
            },
            {
                "name": "Contact Information",
                "score": contact_score,
                "reason": "Contact information detected." if has_contact else "No contact section found.",
                "suggestion": "Add contact information section with links to professional profiles.",
                "missing": [],
            },
        ]

        strengths = []
        improvements = []

        if has_experience:
            strengths.append("Experience section present with detailed entries")
        if has_education:
            strengths.append("Education section present")
        if has_skills:
            strengths.append("Skills section present")
        if has_summary:
            strengths.append("Professional summary provides context")
        if keyword_count >= 5:
            strengths.append(f"Good keyword coverage ({keyword_count} keywords)")
        if len(section_types) >= 5:
            strengths.append("Comprehensive section coverage")

        if not has_experience:
            improvements.append("Add work experience section")
        if not has_education:
            improvements.append("Add education section")
        if not has_skills:
            improvements.append("Add skills section")
        if keyword_count < 5:
            improvements.append("Increase keyword coverage with relevant technologies")
        if not has_summary:
            improvements.append("Add professional summary for better ATS parsing")

        return {
            "overall": overall,
            "categories": categories,
            "strengths": strengths,
            "improvements": improvements,
        }

    async def analyze_health(self, resume_id: uuid.UUID, user_id: uuid.UUID) -> dict:
        resume = await self.resume_repo.get_with_sections(resume_id)
        if not resume or resume.user_id != user_id:
            raise NotFoundError("Resume not found.")
        sections = resume.sections or []
        text = self._extract_text(sections)

        section_types = {s.section_type for s in sections}

        score = 50
        strengths = []
        improvements = []
        recommendations = []

        if "experience" in section_types:
            score += 15
            strengths.append({"label": "Experience", "detail": "Work experience section present"})
        else:
            improvements.append({"label": "Experience", "detail": "Work experience section missing"})
            recommendations.append("Add detailed work experience with achievements")

        if "education" in section_types:
            score += 10
            strengths.append({"label": "Education", "detail": "Education section present"})
        else:
            improvements.append({"label": "Education", "detail": "Education section missing"})
            recommendations.append("Add educational background")

        if "skills" in section_types:
            score += 10
            keywords = self._count_keywords(text)
            if len(keywords) >= 5:
                strengths.append({"label": "Technical Skills", "detail": f"Strong technical skills ({len(keywords)} keywords)"})
                score += 5
            else:
                strengths.append({"label": "Skills", "detail": "Skills section present"})
        else:
            improvements.append({"label": "Skills", "detail": "Skills section missing"})
            recommendations.append("Add a dedicated skills section")

        if "summary" in section_types:
            summary_text = ""
            for s in sections:
                if s.section_type == "summary":
                    content = s.content or {}
                    summary_text = content.get("text", "") if isinstance(content, dict) else str(content)
            if len(summary_text) > 100:
                strengths.append({"label": "Summary", "detail": "Well-written professional summary"})
                score += 5
            else:
                improvements.append({"label": "Summary", "detail": "Summary too generic or brief"})
                recommendations.append("Expand your professional summary with specific achievements")
        else:
            improvements.append({"label": "Summary", "detail": "Professional summary missing"})
            recommendations.append("Add a professional summary")

        if "projects" in section_types:
            strengths.append({"label": "Projects", "detail": "Projects section showcases practical work"})
            score += 5

        word_count = len(text.split())
        if word_count > 300:
            strengths.append({"label": "Content", "detail": f"Excellent content length ({word_count} words)"})
            score += 5
        elif word_count < 100:
            improvements.append({"label": "Content", "detail": "Content too brief"})
            recommendations.append("Expand resume content with more detail")

        has_quantified = any(c.isdigit() for c in text) and any(c in text for c in ["%", "x", "K", "M"])
        if has_quantified:
            strengths.append({"label": "Achievements", "detail": "Includes quantified achievements"})
            score += 5
        else:
            improvements.append({"label": "Achievements", "detail": "Missing quantified achievements"})
            recommendations.append("Add measurable achievements with numbers and percentages")

        ats_analysis = await self.analyze_ats(resume_id, user_id, None)
        ats_score = ats_analysis["overall"]
        if ats_score >= 70:
            strengths.append({"label": "ATS Compatibility", "detail": f"Good ATS score ({ats_score}%)"})
            score += 5
        elif ats_score < 50:
            improvements.append({"label": "ATS Compatibility", "detail": f"Low ATS score ({ats_score}%)"})
            recommendations.append("Improve keyword coverage for better ATS matching")

        overall = min(100, score)

        return {
            "overall": overall,
            "strengths": strengths,
            "improvements": improvements,
            "recommendations": list(dict.fromkeys(recommendations)),
        }

    async def analyze_resume(self, resume_id: uuid.UUID, user_id: uuid.UUID, job_id: uuid.UUID | None = None) -> dict:
        ats = await self.analyze_ats(resume_id, user_id, job_id)
        health = await self.analyze_health(resume_id, user_id)
        return {"ats": ats, "health": health}

    async def export_resume_pdf(self, resume_id: uuid.UUID, user_id: uuid.UUID) -> bytes:
        resume = await self.resume_repo.get_with_sections(resume_id)
        if not resume or resume.user_id != user_id:
            raise NotFoundError("Resume not found.")
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
            import io

            buf = io.BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=letter,
                                     topMargin=0.75*inch, bottomMargin=0.75*inch,
                                     leftMargin=0.75*inch, rightMargin=0.75*inch)
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle("ResumeTitle", parent=styles["Title"],
                                          fontSize=18, spaceAfter=6, alignment=1)
            normal = styles["Normal"]
            section_style = ParagraphStyle("SectionHeading", parent=styles["Heading2"],
                                            fontSize=13, spaceBefore=12, spaceAfter=4,
                                            textColor="#2563eb")

            elements = []
            elements.append(Paragraph(resume.title or "Resume", title_style))
            if resume.template:
                elements.append(Paragraph(f"Template: {resume.template}", normal))
            elements.append(Spacer(1, 12))

            for section in (resume.sections or []):
                elements.append(Paragraph(section.title or section.section_type, section_style))
                content = section.content or {}
                text = content.get("text", "") if isinstance(content, dict) else str(content)
                for line in text.split("\n"):
                    if line.strip():
                        elements.append(Paragraph(line.strip(), normal))
                elements.append(Spacer(1, 8))

            doc.build(elements)
            return buf.getvalue()
        except ImportError:
            raise ImportError("reportlab is required for PDF export")

    async def export_resume_docx(self, resume_id: uuid.UUID, user_id: uuid.UUID) -> bytes:
        resume = await self.resume_repo.get_with_sections(resume_id)
        if not resume or resume.user_id != user_id:
            raise NotFoundError("Resume not found.")
        try:
            from docx import Document
            import io

            buf = io.BytesIO()
            doc = Document()

            doc.add_heading(resume.title or "Resume", 0)
            if resume.template:
                doc.add_paragraph(f"Template: {resume.template}")

            for section in (resume.sections or []):
                doc.add_heading(section.title or section.section_type, level=1)
                content = section.content or {}
                text = content.get("text", "") if isinstance(content, dict) else str(content)
                for line in text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip())

            doc.save(buf)
            return buf.getvalue()
        except ImportError:
            raise ImportError("python-docx is required for DOCX export")
