import uuid

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.exceptions import NotFoundError
from app.models import CoverLetter
from app.repositories import CoverLetterRepository
from app.services.audit import AuditService


class CoverLetterService:
    def __init__(self, session: AsyncSession):
        self.session = session
        self.repo = CoverLetterRepository(session)
        self.audit_service = AuditService(session)

    # ── CRUD ──

    async def create(self, user_id: uuid.UUID, data: dict) -> CoverLetter:
        cl = CoverLetter(
            user_id=user_id,
            job_id=uuid.UUID(data["job_id"]) if data.get("job_id") else None,
            resume_id=uuid.UUID(data["resume_id"]) if data.get("resume_id") else None,
            title=data.get("title"),
            company_name=data.get("company_name"),
            job_title=data.get("job_title"),
            hiring_manager=data.get("hiring_manager"),
            template=data.get("template"),
            tone=data.get("tone"),
            content=data.get("content"),
        )
        created = await self.repo.create(cl)
        await self.audit_service.log("COVER_LETTER_CREATED", user_id=user_id, entity="cover_letter", entity_id=created.id, outcome="success")
        return await self.get(created.id, user_id)

    async def get(self, cover_letter_id: uuid.UUID, user_id: uuid.UUID) -> CoverLetter:
        cl = await self.repo.get_by_id(cover_letter_id)
        if not cl or cl.user_id != user_id:
            raise NotFoundError("Cover letter not found.")
        return cl

    async def update(self, cover_letter_id: uuid.UUID, user_id: uuid.UUID, data: dict) -> CoverLetter:
        cl = await self.get(cover_letter_id, user_id)
        for key, value in data.items():
            if value is not None and hasattr(cl, key):
                setattr(cl, key, value)
            elif key == "hiring_manager":
                cl.hiring_manager = value
        await self.repo.update(cl)
        await self.audit_service.log("COVER_LETTER_UPDATED", user_id=user_id, entity="cover_letter", entity_id=cover_letter_id, outcome="success")
        return cl

    async def delete(self, cover_letter_id: uuid.UUID, user_id: uuid.UUID) -> None:
        cl = await self.get(cover_letter_id, user_id)
        await self.repo.delete(cl)
        await self.audit_service.log("COVER_LETTER_DELETED", user_id=user_id, entity="cover_letter", entity_id=cover_letter_id, outcome="success")

    async def list_by_user(self, user_id: uuid.UUID, status: str | None = None) -> list[CoverLetter]:
        stmt = select(CoverLetter).where(CoverLetter.user_id == user_id)
        if status:
            stmt = stmt.where(CoverLetter.status == status)
        stmt = stmt.order_by(CoverLetter.updated_at.desc())
        result = await self.session.execute(stmt)
        return result.scalars().all()

    async def duplicate(self, cover_letter_id: uuid.UUID, user_id: uuid.UUID) -> CoverLetter:
        cl = await self.get(cover_letter_id, user_id)
        new_cl = CoverLetter(
            user_id=user_id,
            job_id=cl.job_id,
            resume_id=cl.resume_id,
            title=f"{cl.title or 'Cover Letter'} (Copy)",
            company_name=cl.company_name,
            job_title=cl.job_title,
            template=cl.template,
            tone=cl.tone,
            content=cl.content,
        )
        created = await self.repo.create(new_cl)
        await self.audit_service.log("COVER_LETTER_DUPLICATED", user_id=user_id, entity="cover_letter", entity_id=created.id, outcome="success")
        return created

    # ── AI Generation ──

    async def generate_ai(
        self,
        user_id: uuid.UUID,
        job_id: uuid.UUID,
        resume_id: uuid.UUID,
        tone: str = "professional",
        template: str = "modern",
        hiring_manager: str | None = None,
        additional_notes: str | None = None,
    ) -> CoverLetter:
        from app.ai.features.cover_letter import ai_generate_cover_letter
        from app.services.job import JobService
        from app.services.resume import ResumeService

        job_svc = JobService(self.session)
        resume_svc = ResumeService(self.session)

        job = await job_svc.get_job(job_id)
        resume = await resume_svc.get_resume(resume_id, user_id)
        resume_text = resume_svc._extract_text(resume.sections or [])

        job_title = getattr(job, "title", "") or ""
        company_name = getattr(job, "company", "") or ""
        job_desc = getattr(job, "description", "") or ""

        result = await ai_generate_cover_letter(
            job_title=job_title,
            company_name=company_name,
            job_description=job_desc,
            resume_text=resume_text,
            tone=tone,
            style=template,
            hiring_manager=hiring_manager or "",
        )

        content = result.get("cover_letter", "")

        cl = CoverLetter(
            user_id=user_id,
            job_id=job_id,
            resume_id=resume_id,
            title=f"Cover Letter - {job_title or company_name or 'Generated'}",
            company_name=company_name,
            job_title=job_title,
            hiring_manager=hiring_manager,
            template=template,
            tone=tone,
            content=content,
        )
        created = await self.repo.create(cl)
        await self.audit_service.log("COVER_LETTER_GENERATED_AI", user_id=user_id, entity="cover_letter", entity_id=created.id, outcome="success")
        return created

    # ── AI-Assisted Editing ──

    async def ai_assist(self, content: str, instruction: str, context: str | None = None) -> str:
        from app.ai.features.cover_letter import ai_assist_cover_letter

        result = await ai_assist_cover_letter(
            content=content,
            instruction=instruction,
            context=context or "",
        )
        return result.get("edited_content", content)

    # ── Export ──

    async def export_pdf(self, cover_letter_id: uuid.UUID, user_id: uuid.UUID) -> bytes:
        cl = await self.get(cover_letter_id, user_id)
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
            import io

            buf = io.BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=1*inch, bottomMargin=1*inch,
                                     leftMargin=1*inch, rightMargin=1*inch)
            title_style = ParagraphStyle("Title", fontSize=16, spaceAfter=12)
            normal = ParagraphStyle("Normal", fontSize=11, leading=16, spaceAfter=6)

            elements = []
            elements.append(Paragraph(cl.title or "Cover Letter", title_style))
            if cl.company_name:
                elements.append(Paragraph(f"Company: {cl.company_name}", normal))
            if cl.job_title:
                elements.append(Paragraph(f"Position: {cl.job_title}", normal))
            elements.append(Spacer(1, 12))

            for line in (cl.content or "").split("\n"):
                if line.strip():
                    elements.append(Paragraph(line.strip(), normal))
                else:
                    elements.append(Spacer(1, 6))

            doc.build(elements)
            return buf.getvalue()
        except ImportError:
            raise ImportError("reportlab is required for PDF export")

    async def export_docx(self, cover_letter_id: uuid.UUID, user_id: uuid.UUID) -> bytes:
        cl = await self.get(cover_letter_id, user_id)
        try:
            import io
            from docx import Document
            doc = Document()
            doc.add_heading(cl.title or "Cover Letter", 0)
            if cl.company_name:
                doc.add_paragraph(f"Company: {cl.company_name}")
            if cl.job_title:
                doc.add_paragraph(f"Position: {cl.job_title}")
            for line in (cl.content or "").split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except ImportError:
            raise ImportError("python-docx is required for DOCX export")

    # ── Application Package ──

    async def create_application_package(
        self,
        user_id: uuid.UUID,
        resume_id: uuid.UUID,
        cover_letter_id: uuid.UUID,
        job_id: uuid.UUID,
        notes: str | None = None,
    ) -> dict:
        from app.services.resume import ResumeService
        from app.services.job import JobService

        resume_svc = ResumeService(self.session)
        job_svc = JobService(self.session)

        resume = await resume_svc.get_resume(resume_id, user_id)
        cover_letter = await self.get(cover_letter_id, user_id)
        job = await job_svc.get_job(job_id)

        from app.schemas.resume import ResumeResponse
        from app.schemas.cover_letter import CoverLetterResponse

        return {
            "resume": ResumeResponse.model_validate(resume).model_dump(),
            "cover_letter": CoverLetterResponse.model_validate(cover_letter).model_dump(),
            "job": {
                "id": str(job.id) if hasattr(job, "id") else "",
                "title": getattr(job, "title", None),
                "company": getattr(job, "company", None),
                "description": (getattr(job, "description", None) or "")[:500],
            },
            "notes": notes,
        }
