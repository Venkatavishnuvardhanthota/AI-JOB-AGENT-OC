import io
import json
import logging
from pathlib import Path

from app.core.config import settings

logger = logging.getLogger(__name__)


class ResumeGeneratorService:
    BUILTIN_TEMPLATES = {
        "modern": {
            "name": "Modern",
            "description": "Clean, modern layout with blue accent colors",
            "layout_config": {
                "primary_color": "#2563eb",
                "secondary_color": "#1e40af",
                "font_family": "Helvetica",
                "font_size_body": 10,
                "font_size_header": 16,
                "font_size_section": 13,
                "margin_top": 40,
                "margin_bottom": 40,
                "margin_left": 40,
                "margin_right": 40,
                "show_icons": False,
                "section_order": [
                    "summary", "experience", "education", "skills",
                    "projects", "certifications", "languages", "portfolio",
                ],
            },
        },
        "classic": {
            "name": "Classic",
            "description": "Traditional resume layout with serif font",
            "layout_config": {
                "primary_color": "#1f2937",
                "secondary_color": "#4b5563",
                "font_family": "Times-Roman",
                "font_size_body": 11,
                "font_size_header": 18,
                "font_size_section": 14,
                "margin_top": 36,
                "margin_bottom": 36,
                "margin_left": 36,
                "margin_right": 36,
                "show_icons": False,
                "section_order": [
                    "summary", "experience", "education", "skills",
                    "projects", "certifications", "languages", "portfolio",
                ],
            },
        },
        "minimal": {
            "name": "Minimal",
            "description": "Minimal, clean design with lots of white space",
            "layout_config": {
                "primary_color": "#000000",
                "secondary_color": "#374151",
                "font_family": "Helvetica",
                "font_size_body": 9,
                "font_size_header": 14,
                "font_size_section": 11,
                "margin_top": 50,
                "margin_bottom": 50,
                "margin_left": 50,
                "margin_right": 50,
                "show_icons": False,
                "section_order": [
                    "summary", "experience", "education", "skills",
                    "projects", "certifications", "languages", "portfolio",
                ],
            },
        },
    }

    def __init__(self) -> None:
        self.template_dir = Path(settings.RESUME_TEMPLATE_DIR)
        self.template_dir.mkdir(parents=True, exist_ok=True)

    def get_template_config(self, template_name: str) -> dict:
        if template_name in self.BUILTIN_TEMPLATES:
            return self.BUILTIN_TEMPLATES[template_name]
        custom_path = self.template_dir / f"{template_name}.json"
        if custom_path.exists():
            with open(custom_path) as f:
                return json.load(f)
        logger.warning("Template '%s' not found, falling back to 'modern'", template_name)
        return self.BUILTIN_TEMPLATES["modern"]

    def generate_docx(self, snapshot: dict, template_name: str = "modern") -> bytes:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        config = self.get_template_config(template_name).get("layout_config", {})
        doc = Document()

        section = doc.sections[0]
        section.top_margin = Inches(config.get("margin_top", 40) / 25.4)
        section.bottom_margin = Inches(config.get("margin_bottom", 40) / 25.4)
        section.left_margin = Inches(config.get("margin_left", 40) / 25.4)
        section.right_margin = Inches(config.get("margin_right", 40) / 25.4)

        primary_color = config.get("primary_color", "#2563eb")
        font_family = config.get("font_family", "Helvetica")
        font_size_header = config.get("font_size_header", 16)
        font_size_section = config.get("font_size_section", 13)
        font_size_body = config.get("font_size_body", 10)

        def _add_styled_paragraph(text: str, size: int, bold: bool = False, color: str | None = None, alignment: int | None = None):  # noqa: E501
            p = doc.add_paragraph()
            if alignment is not None:
                p.alignment = alignment
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.name = font_family
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*self._hex_to_rgb(color))
            return p

        def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
            hex_color = hex_color.lstrip("#")
            return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

        profile = snapshot.get("profile", {})
        full_name = snapshot.get("full_name") or profile.get("full_name", "")
        headline = profile.get("headline", "")
        email = profile.get("email", "") or snapshot.get("email", "")
        phone = profile.get("phone", "") or snapshot.get("phone", "")
        location = profile.get("location", "")
        linkedin = profile.get("linkedin_url", "") or snapshot.get("linkedin_url", "")
        github = profile.get("github_url", "") or snapshot.get("github_url", "")
        portfolio_url = profile.get("portfolio_url", "")
        summary_text = profile.get("bio", "") or snapshot.get("summary", "")

        # Header
        if full_name:
            _add_styled_paragraph(full_name, font_size_header, bold=True, color=primary_color, alignment=WD_ALIGN_PARAGRAPH.CENTER)  # noqa: E501

        if headline:
            _add_styled_paragraph(headline, 11, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        contact_parts = [p for p in [phone, email, location, linkedin, github, portfolio_url] if p]
        if contact_parts:
            _add_styled_paragraph(" | ".join(contact_parts), 9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        section_order = config.get("section_order", [])

        # Summary
        if "summary" in section_order and summary_text:
            doc.add_paragraph()
            _add_styled_paragraph("Professional Summary", font_size_section, bold=True, color=primary_color)
            _add_styled_paragraph(summary_text, font_size_body)

        # Experience
        if "experience" in section_order:
            experiences = snapshot.get("experience", [])
            if experiences:
                doc.add_paragraph()
                _add_styled_paragraph("Experience", font_size_section, bold=True, color=primary_color)
                for exp in experiences:
                    p = doc.add_paragraph()
                    run = p.add_run(exp.get("title", ""))
                    run.bold = True
                    run.font.size = Pt(font_size_body + 1)
                    run.font.name = font_family
                    run2 = p.add_run(f" at {exp.get('company', '')}")
                    run2.font.size = Pt(font_size_body)
                    run2.font.name = font_family

                    if exp.get("start_date") or exp.get("end_date"):
                        date_str = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present') if exp.get('is_current') else exp.get('end_date', '')}"  # noqa: E501
                        _add_styled_paragraph(date_str, font_size_body - 1, color="#6b7280")

                    if exp.get("description"):
                        _add_styled_paragraph(exp["description"][:500], font_size_body)

        # Education
        if "education" in section_order:
            educations = snapshot.get("education", [])
            if educations:
                doc.add_paragraph()
                _add_styled_paragraph("Education", font_size_section, bold=True, color=primary_color)
                for edu in educations:
                    p = doc.add_paragraph()
                    run = p.add_run(edu.get("institution", ""))
                    run.bold = True
                    run.font.size = Pt(font_size_body + 1)
                    run.font.name = font_family
                    _add_styled_paragraph(f"{edu.get('degree', '')}{', ' + edu.get('field_of_study', '') if edu.get('field_of_study') else ''}", font_size_body)  # noqa: E501
                    if edu.get("gpa"):
                        _add_styled_paragraph(f"GPA: {edu['gpa']}", font_size_body)

        # Skills
        if "skills" in section_order:
            skills = snapshot.get("skills", [])
            if skills:
                doc.add_paragraph()
                _add_styled_paragraph("Skills", font_size_section, bold=True, color=primary_color)
                skill_text = ", ".join(s.get("name", "") for s in skills)
                _add_styled_paragraph(skill_text, font_size_body)

        # Projects
        if "projects" in section_order:
            projects = snapshot.get("projects", [])
            if projects:
                doc.add_paragraph()
                _add_styled_paragraph("Projects", font_size_section, bold=True, color=primary_color)
                for proj in projects:
                    p = doc.add_paragraph()
                    run = p.add_run(proj.get("name", ""))
                    run.bold = True
                    run.font.size = Pt(font_size_body + 1)
                    run.font.name = font_family
                    if proj.get("description"):
                        _add_styled_paragraph(proj["description"][:500], font_size_body)

        # Certifications
        if "certifications" in section_order:
            certs = snapshot.get("certifications", [])
            if certs:
                doc.add_paragraph()
                _add_styled_paragraph("Certifications", font_size_section, bold=True, color=primary_color)
                for cert in certs:
                    p = doc.add_paragraph()
                    run = p.add_run(cert.get("name", ""))
                    run.font.size = Pt(font_size_body)
                    run.font.name = font_family
                    if cert.get("issuer"):
                        run2 = p.add_run(f" - {cert['issuer']}")
                        run2.font.size = Pt(font_size_body)
                        run2.font.name = font_family

        # Languages
        if "languages" in section_order:
            languages = snapshot.get("languages", [])
            if languages:
                doc.add_paragraph()
                _add_styled_paragraph("Languages", font_size_section, bold=True, color=primary_color)
                lang_text = ", ".join(
                    f"{lang.get('name', '')} ({lang.get('proficiency', '')})" if lang.get('proficiency') else lang.get('name', '')  # noqa: E501
                    for lang in languages
                )
                _add_styled_paragraph(lang_text, font_size_body)

        # Portfolio
        if "portfolio" in section_order:
            portfolio = snapshot.get("portfolio_items", [])
            if portfolio:
                doc.add_paragraph()
                _add_styled_paragraph("Portfolio", font_size_section, bold=True, color=primary_color)
                for item in portfolio:
                    p = doc.add_paragraph()
                    run = p.add_run(item.get("title", ""))
                    run.bold = True
                    run.font.size = Pt(font_size_body + 1)
                    run.font.name = font_family
                    if item.get("description"):
                        _add_styled_paragraph(item["description"][:300], font_size_body)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    def generate_pdf(self, snapshot: dict, template_name: str = "modern") -> bytes:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )

        config = self.get_template_config(template_name).get("layout_config", {})
        primary_color_hex = config.get("primary_color", "#2563eb")
        primary_rgb = self._hex_to_rgb(primary_color_hex)
        primary = colors.Color(primary_rgb[0]/255, primary_rgb[1]/255, primary_rgb[2]/255)
        font_family = config.get("font_family", "Helvetica")
        fs_header = config.get("font_size_header", 16)
        fs_section = config.get("font_size_section", 13)
        fs_body = config.get("font_size_body", 10)

        buf = io.BytesIO()
        config.get("margin_left", 40) * mm / 25.4  # convert mm to points via inches
        doc = SimpleDocTemplate(
            buf,
            pagesize=letter,
            topMargin=config.get("margin_top", 40),
            bottomMargin=config.get("margin_bottom", 40),
            leftMargin=config.get("margin_left", 40),
            rightMargin=config.get("margin_right", 40),
        )

        getSampleStyleSheet()

        style_header = ParagraphStyle(
            "ResumeHeader", fontName=font_family, fontSize=fs_header,
            textColor=primary, alignment=TA_CENTER, spaceAfter=4,
        )
        style_subheader = ParagraphStyle(
            "ResumeSubheader", fontName=font_family, fontSize=11,
            alignment=TA_CENTER, spaceAfter=2,
        )
        style_contact = ParagraphStyle(
            "ResumeContact", fontName=font_family, fontSize=9,
            alignment=TA_CENTER, spaceAfter=10,
        )
        style_section = ParagraphStyle(
            "ResumeSection", fontName=font_family, fontSize=fs_section,
            textColor=primary, spaceBefore=12, spaceAfter=6, leading=fs_section + 4,
        )
        style_body = ParagraphStyle(
            "ResumeBody", fontName=font_family, fontSize=fs_body,
            spaceAfter=4, leading=fs_body + 4,
        )
        style_bold = ParagraphStyle(
            "ResumeBold", fontName=font_family, fontSize=fs_body + 1,
            spaceAfter=2, leading=fs_body + 6,
        )
        style_date = ParagraphStyle(
            "ResumeDate", fontName=font_family, fontSize=fs_body - 1,
            textColor=colors.gray, spaceAfter=2,
        )

        elements: list = []
        profile = snapshot.get("profile", {})
        full_name = snapshot.get("full_name") or profile.get("full_name", "")
        headline = profile.get("headline", "")
        email = profile.get("email", "") or snapshot.get("email", "")
        phone = profile.get("phone", "") or snapshot.get("phone", "")
        location = profile.get("location", "")
        linkedin = profile.get("linkedin_url", "") or snapshot.get("linkedin_url", "")
        github = profile.get("github_url", "") or snapshot.get("github_url", "")
        portfolio_url = profile.get("portfolio_url", "")
        summary_text = profile.get("bio", "") or snapshot.get("summary", "")

        if full_name:
            elements.append(Paragraph(full_name, style_header))
        if headline:
            elements.append(Paragraph(headline, style_subheader))
        contact_parts = [p for p in [phone, email, location, linkedin, github, portfolio_url] if p]
        if contact_parts:
            elements.append(Paragraph(" | ".join(contact_parts), style_contact))

        section_order = config.get("section_order", [])

        if "summary" in section_order and summary_text:
            elements.append(Spacer(1, 6))
            elements.append(Paragraph("Professional Summary", style_section))
            elements.append(Paragraph(summary_text, style_body))

        if "experience" in section_order:
            experiences = snapshot.get("experience", [])
            if experiences:
                elements.append(Spacer(1, 6))
                elements.append(Paragraph("Experience", style_section))
                for exp in experiences:
                    elements.append(Paragraph(
                        f"<b>{exp.get('title', '')}</b> at <b>{exp.get('company', '')}</b>",
                        style_bold,
                    ))
                    if exp.get("start_date") or exp.get("end_date"):
                        date_str = f"{exp.get('start_date', '')} - {'Present' if exp.get('is_current') else (exp.get('end_date', '') or '')}"  # noqa: E501
                        elements.append(Paragraph(date_str, style_date))
                    if exp.get("description"):
                        elements.append(Paragraph(exp["description"][:500], style_body))

        if "education" in section_order:
            educations = snapshot.get("education", [])
            if educations:
                elements.append(Spacer(1, 6))
                elements.append(Paragraph("Education", style_section))
                for edu in educations:
                    elements.append(Paragraph(
                        f"<b>{edu.get('institution', '')}</b>", style_bold,
                    ))
                    degree_parts = [edu.get("degree", "")]
                    if edu.get("field_of_study"):
                        degree_parts.append(edu["field_of_study"])
                    elements.append(Paragraph(", ".join(degree_parts), style_body))
                    if edu.get("gpa"):
                        elements.append(Paragraph(f"GPA: {edu['gpa']}", style_body))

        if "skills" in section_order:
            skills = snapshot.get("skills", [])
            if skills:
                elements.append(Spacer(1, 6))
                elements.append(Paragraph("Skills", style_section))
                skill_text = ", ".join(s.get("name", "") for s in skills)
                elements.append(Paragraph(skill_text, style_body))

        if "projects" in section_order:
            projects = snapshot.get("projects", [])
            if projects:
                elements.append(Spacer(1, 6))
                elements.append(Paragraph("Projects", style_section))
                for proj in projects:
                    elements.append(Paragraph(
                        f"<b>{proj.get('name', '')}</b>", style_bold,
                    ))
                    if proj.get("description"):
                        elements.append(Paragraph(proj["description"][:500], style_body))

        if "certifications" in section_order:
            certs = snapshot.get("certifications", [])
            if certs:
                elements.append(Spacer(1, 6))
                elements.append(Paragraph("Certifications", style_section))
                for cert in certs:
                    cert_text = cert.get("name", "")
                    if cert.get("issuer"):
                        cert_text += f" - {cert['issuer']}"
                    elements.append(Paragraph(cert_text, style_body))

        if "languages" in section_order:
            languages = snapshot.get("languages", [])
            if languages:
                elements.append(Spacer(1, 6))
                elements.append(Paragraph("Languages", style_section))
                lang_text = ", ".join(
                    f"{lang.get('name', '')} ({lang.get('proficiency', '')})" if lang.get('proficiency') else lang.get('name', '')  # noqa: E501
                    for lang in languages
                )
                elements.append(Paragraph(lang_text, style_body))

        if "portfolio" in section_order:
            portfolio = snapshot.get("portfolio_items", [])
            if portfolio:
                elements.append(Spacer(1, 6))
                elements.append(Paragraph("Portfolio", style_section))
                for item in portfolio:
                    elements.append(Paragraph(
                        f"<b>{item.get('title', '')}</b>", style_bold,
                    ))
                    if item.get("description"):
                        elements.append(Paragraph(item["description"][:300], style_body))

        doc.build(elements)
        buf.seek(0)
        return buf.getvalue()

    def _hex_to_rgb(self, hex_color: str) -> tuple[int, int, int]:
        hex_color = hex_color.lstrip("#")
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

    def get_available_templates(self) -> list[dict]:
        templates = []
        for key, tmpl in self.BUILTIN_TEMPLATES.items():
            templates.append({"id": key, "name": tmpl["name"], "description": tmpl["description"], "is_system": True})
        if self.template_dir.exists():
            for f in self.template_dir.glob("*.json"):
                try:
                    with open(f) as fh:
                        data = json.load(fh)
                    templates.append({
                        "id": f.stem,
                        "name": data.get("name", f.stem),
                        "description": data.get("description", ""),
                        "is_system": False,
                    })
                except (json.JSONDecodeError, OSError):
                    pass
        return templates
